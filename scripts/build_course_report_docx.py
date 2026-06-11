from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
DRAFT = ROOT / "COURSE_REPORT_DRAFT.md"
OUT = ROOT / "预言机不确定性下DeFi抵押借贷清算机制的风险权衡与改进研究_结课报告.docx"


TITLE = "预言机不确定性下 DeFi 抵押借贷清算机制的风险权衡与改进研究"


PROBLEM_MODELING = """
## 三、问题建模与指标定义

### 3.1 基本账户模型
本文采用简化的单资产抵押模型。借款人抵押 `C` 个 ETH，借出 `D` 个 USDC。市场参考价格记为 `P_market,t`，预言机价格记为 `P_oracle,t`。由于预言机延迟或异常，二者可能不完全一致。

真实健康因子定义为：

```text
HF_true,t = C_t * P_market,t * LT / D_t
```

预言机健康因子定义为：

```text
HF_oracle,t = C_t * P_oracle,t * LT / D_t
```

真实健康因子用于实验评估账户在市场参考价格下是否安全；预言机健康因子用于模拟智能合约在链上能够看到的账户状态。真实市场价格并不直接参与链上清算决策，而是作为评估指标的参考基准。

### 3.2 两类风险
本文关注两类风险。第一类是协议侧坏账风险。当账户剩余债务大于剩余抵押品在市场价格下的价值时，协议产生坏账：

```text
BadDebt_t = max(0, D_t - C_t * P_market,t)
```

第二类是用户侧误清算风险。当账户在市场参考价格下仍然安全，即 `HF_true,t >= 1`，但由于预言机价格触发清算，用户承担的清算折扣损失被记为误清算损失。若清算偿还金额为 `R_t`，清算奖励比例为 `bonus`，则该次清算给用户造成的折扣损失近似为：

```text
UserLoss_t = R_t * bonus
```

若该次清算发生时 `HF_true,t >= 1`，则该损失计入误清算损失。坏账是协议侧风险，用户损失是借款人侧风险，二者属于不同主体的风险指标，不能简单相加。

### 3.3 研究定位
本文不是提出可直接部署于主网的完整 DeFi 协议，而是构建一个简化实验原型，分析预言机价格不确定时不同清算机制在协议坏账风险与用户误清算风险之间的权衡。Aave 和 Compound 等协议的文档均表明，健康因子、抵押因子、清算阈值和清算比例是抵押借贷协议控制风险的重要参数[1][2][3][4]；Chainlink 文档则说明，链上合约需要通过预言机读取链下资产价格[5][6]。因此，本文的机制设计重点不是重新定义抵押借贷协议，而是研究当预言机输入不确定时，清算强度如何根据风险状态进行可解释调节。
"""


IMPLEMENTATION = """
## 五、原型实现与可复现性

### 5.1 Python 仿真实现
本文实现了一个轻量级 Python 仿真系统，用于生成价格路径、模拟预言机延迟、计算健康因子区间，并比较不同清算机制。核心仿真代码位于 `src/uspl/simulator.py`，命令行实验脚本位于 `src/uspl/run_demo.py`。运行方式为：

```bash
python3 src/uspl/run_demo.py
```

该脚本输出 `outputs/demo_metrics.csv`、`outputs/demo_paths.png` 和 `outputs/mechanism_metrics.png`。其中，`demo_metrics.csv` 保存不同机制在各场景下的指标结果；`demo_paths.png` 展示市场价格、预言机价格、价格区间和健康因子区间；`mechanism_metrics.png` 展示不同机制在坏账、误清算损失、用户损失和清算延迟上的对比。

参数敏感性分析代码位于 `src/uspl/sensitivity_analysis.py`，运行方式为：

```bash
python3 src/uspl/sensitivity_analysis.py
```

该脚本输出 `outputs/uspl_gamma_sensitivity.csv` 和 `outputs/uspl_gamma_sensitivity.png`，用于分析 `gamma` 参数对清算强度、误清算损失和坏账的影响。

### 5.2 Streamlit 展示
项目还实现了一个 Streamlit 展示界面 `app/streamlit_app.py`。该界面支持调整价格场景、预言机延迟、区间宽度、`gamma`、`cap_min` 和 `cap_max` 等参数，并可视化价格路径、健康因子区间和清算事件。它主要用于交互式展示，帮助说明 USPL 机制在不同场景下的行为。

### 5.3 Solidity 合约草案
本文给出 Solidity 合约草案以说明 USPL 的链上可表达性。合约文件包括 `contracts/MockOracle.sol` 和 `contracts/SimpleLendingUSPL.sol`。其中，`MockOracle.sol` 用于模拟预言机价格更新；`SimpleLendingUSPL.sol` 实现抵押、借款、健康因子计算、价格区间、分区判断和清算函数。

需要说明的是，Python 仿真中的不确定性区间是动态计算的，而 Solidity 草案为了保持链上逻辑简洁，采用固定 `uncertaintyWidthWad` 表示价格区间宽度。也就是说，合约草案实现的是链上可执行的简化版本，而 Python 仿真承担主要实验分析功能。后续若要进一步工程化，可以将动态不确定性宽度设计为由链上波动统计、预言机更新时间差或外部风险参数共同决定。
"""


ABSTRACT = """DeFi 抵押借贷协议通常依赖智能合约和预言机价格实现自动化清算[1][2][5]。当抵押品价格下跌导致账户健康因子低于清算边界时，清算人可以偿还债务并获得折价抵押品，从而维护协议偿付能力。然而，清算决策高度依赖预言机价格输入：价格更新延迟可能导致清算过晚并产生坏账，短时异常价格则可能导致账户在真实市场价格下仍然安全却被错误清算。本文关注的问题是：当预言机价格存在不确定性时，DeFi 借贷协议如何在协议偿付保护和借款人保护之间进行更稳健的清算强度调节。

为分析这一问题，本文构建了一个简化 DeFi 抵押借贷实验原型，并提出 USPL（Uncertainty-Scaled Partial Liquidation）机制。该机制包含三个可解释的传导环节：首先，根据预言机价格与短期 TWAP 的偏离以及近期波动构造动态不确定性宽度；其次，将单点预言机价格扩展为价格区间，并进一步计算健康因子上下界；最后，在健康因子区间跨越清算边界的状态下，根据账户不安全概率、最低偿付保护需求和误清算预算自适应确定 close factor，使清算动作从二元触发转变为强度可调的风险控制规则。

本文实验采用真实 ETH/USD 价格回放，并采用 safety-first ε-constraint 思路进行评价：借鉴 Roy 的 safety-first principle，先控制跌破最低可接受安全阈值的灾难风险；再借鉴多目标优化中的 ε-constraint method，将协议坏账和尾部坏账作为偿付安全约束，而不是与用户损失和执行频率任意加权；最后在满足偿付安全约束的机制集合内比较误清算损失、用户损失、执行频率，并使用 TOPSIS 作为辅助排序。真实数据实验表明，TWAP 在反事实 oracle shock 中未触发误清算，但在代表性真实下跌窗口中产生最大坏账 336.71、坏账 ES95 为 323.43。20 个最大回撤滚动窗口仅作为经验压力样本，结果显示 TWAP 的坏账窗口数为 20/20，USPL 为 5/20，低于 buffer 的 10/20 和 fixed 的 13/20。与此同时，USPL 将反事实 oracle shock 场景下的误清算损失由 fixed/buffer 的 35.53 降至 24.15。本文将 USPL 定位为一种将预言机不确定性显式传导到清算比例的可解释机制，用于分析 DeFi 清算机制中的风险权衡。"""


DIFFERENCE_SECTION = """
### 3.5 本文改进点与已有机制的区别
需要明确的是，安全区、不确定区和清算区的状态划分并不是本文唯一或主要的创新来源。已有 DeFi 协议和风险控制方法中已经存在健康因子阈值、安全缓冲、价格过滤和风险分层等思想。本文的改进重点在于把这些思想连接成一条可计算、可解释的清算强度传导链条。

第一，本文没有只使用单点预言机价格计算健康因子，而是根据预言机价格相对 TWAP 的偏离和近期波动构造动态不确定性宽度。这样可以区分普通价格波动和预言机输入本身不稳定的状态。

第二，本文没有把价格过滤作为唯一处理方式。TWAP 可以降低短时异常价格影响，但在持续下跌场景中可能反应滞后。因此，本文保留即时预言机价格的信息，同时通过价格区间和健康因子区间表达不确定性，而不是简单用平滑价格替代当前价格。

第三，本文没有把边界状态简单处理为“可清算”或“不可清算”。当健康因子区间跨越清算边界时，账户风险具有不确定性。USPL 将该不确定性映射到最大清算比例，使清算机制从二元阈值判断扩展为清算强度调节问题。

第四，本文通过 fixed、TWAP、buffer 和 USPL 的对照实验，以及 `gamma` 参数敏感性分析，展示不同机制在坏账风险、误清算损失和用户损失之间的权衡。本文的贡献不在于证明某一机制全场景最优，而在于给出一个可以复现实验、可以被合约表达、且便于讨论参数权衡的机制分析框架。
"""


REFERENCES = """
## 参考文献

[1] Aave. Health Factor & Liquidations. https://aave.com/help/borrowing/liquidations

[2] Aave. Aave V3 Overview. https://aave.com/docs/aave-v3/overview

[3] Compound Finance. Comptroller, Compound v2 Documentation. https://docs.compound.finance/v2/comptroller/

[4] Compound Finance. Liquidation, Compound III Documentation. https://docs.compound.finance/liquidation/

[5] Chainlink. Chainlink Data Feeds Documentation. https://docs.chain.link/data-feeds

[6] Chainlink. Consuming Data Feeds. https://docs.chain.link/data-feeds/getting-started

[7] Bank for International Settlements. The oracle problem and the future of DeFi. BIS Bulletin No. 76, 2023. https://www.bis.org/publ/bisbull76.pdf

[8] OWASP. SC02:2025 Price Oracle Manipulation, Smart Contract Top 10. https://owasp.org/www-project-smart-contract-top-10/2025/en/src/SC02-price-oracle-manipulation.html

[9] Arora, S., et al. SecDeLP: Secure Decentralized Lending Platforms against Oracle Manipulation Attacks. University of Oregon Technical Report, 2023. https://www.cs.uoregon.edu/Reports/DRP-202306-Arora.pdf

[10] Qin, K., Zhou, L., Gervais, A. Quantifying Blockchain Extractable Value: How dark is the forest? IEEE Symposium on Security and Privacy, 2022.

[11] Werner, S. M., Perez, D., Gudgeon, L., Klages-Mundt, A., Harz, D., Knottenbelt, W. J. SoK: Decentralized Finance (DeFi). Proceedings of the 4th ACM Conference on Advances in Financial Technologies, 2022.

[12] Roy, A. D. Safety First and the Holding of Assets. Econometrica, 1952.

[13] Rockafellar, R. T., and Uryasev, S. Optimization of conditional value-at-risk. Journal of Risk, 2000.

[14] GAMS. eps-Constraint Method for Multiobjective Optimization. https://forum.gams.com/t/eps-constraint-method-for-multiobjective-optimization/4123

[15] Hwang, C. L., and Yoon, K. Multiple Attribute Decision Making: Methods and Applications. Springer, 1981.

## 附录：项目文件说明

本文使用的实验原型文件结构如下：

```text
src/uspl/simulator.py              核心仿真逻辑
src/uspl/run_demo.py               机制对比实验脚本
src/uspl/sensitivity_analysis.py   参数敏感性分析脚本
app/streamlit_app.py               Streamlit 可视化展示
contracts/MockOracle.sol           模拟预言机合约
contracts/SimpleLendingUSPL.sol    USPL 简化借贷合约
outputs/demo_metrics.csv           机制对比指标
outputs/demo_paths.png             价格路径与健康因子图
outputs/mechanism_metrics.png      机制指标对比图
outputs/uspl_gamma_sensitivity.csv gamma 敏感性指标
outputs/uspl_gamma_sensitivity.png gamma 敏感性图
```

复现实验的主要命令如下：

```bash
python3 src/uspl/run_demo.py
python3 src/uspl/sensitivity_analysis.py
```
"""


def set_run_font(run, size: float = 12, bold: bool = False, font_cn: str = "宋体"):
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_cn)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_body(paragraph, first_line: bool = True):
    fmt = paragraph.paragraph_format
    fmt.line_spacing = 1.25
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    if first_line:
        fmt.first_line_indent = Pt(24)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_text_with_citations(paragraph, text: str, size: float = 12, bold: bool = False, font_cn: str = "宋体"):
    text = text.replace("**", "").replace("`", "")
    pattern = re.compile(r"(\[\d+\](?:\[\d+\])*)")
    pos = 0
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos:match.start()])
            set_run_font(run, size=size, bold=bold, font_cn=font_cn)
        for item in re.findall(r"\[\d+\]", match.group(0)):
            run = paragraph.add_run(item)
            set_run_font(run, size=size, bold=bold, font_cn=font_cn)
            run.font.superscript = True
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        set_run_font(run, size=size, bold=bold, font_cn=font_cn)


def add_centered(paragraph, text: str, size: float, bold: bool = False, font_cn: str = "宋体"):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, font_cn=font_cn)
    return run


def set_cell_text(cell, text: str, size: float = 12, bold: bool = False, align=WD_ALIGN_PARAGRAPH.CENTER):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold)


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def build_cover(doc: Document):
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(3.17)
    sec.right_margin = Cm(3.17)

    p = doc.add_paragraph()
    add_centered(p, "对外经济贸易大学", 22, bold=True, font_cn="黑体")
    p.paragraph_format.space_after = Pt(14)

    p = doc.add_paragraph()
    add_centered(p, "2025 — 2026学年第二学期期末考试", 15, bold=False, font_cn="宋体")
    p.paragraph_format.space_after = Pt(36)

    p = doc.add_paragraph()
    add_centered(p, "论文题目", 14, bold=False, font_cn="宋体")
    p.paragraph_format.space_after = Pt(8)

    p = doc.add_paragraph()
    add_centered(p, TITLE, 16, bold=True, font_cn="黑体")
    p.paragraph_format.space_after = Pt(42)

    fields = [
        ("课程代码及课序号", ""),
        ("课程名称", "区块链技术及应用"),
        ("学    号", ""),
        ("姓    名", ""),
        ("学    院", "统计学院"),
        ("专    业", "数据科学与大数据技术"),
        ("考试时间", ""),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(3.9)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(label + "    ")
        set_run_font(run, 12)
        shown = value + "____________" if value else "____________________________"
        value_run = p.add_run(shown)
        set_run_font(value_run, 12)

    doc.add_paragraph("\n\n")
    p = doc.add_paragraph()
    add_centered(p, "成绩", 14, font_cn="宋体")
    p.paragraph_format.space_before = Pt(28)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("________________")
    set_run_font(run, 12)

    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def prepare_markdown() -> str:
    text = DRAFT.read_text(encoding="utf-8")
    text = re.sub(r"^# .+$", f"# {TITLE}", text, count=1, flags=re.MULTILINE)
    text = re.sub(
        r"## 摘要\n.*?\n\n\*\*关键词\*\*",
        "## 摘要\n" + ABSTRACT.strip() + "\n\n**关键词**",
        text,
        count=1,
        flags=re.DOTALL,
    )
    text = text.replace(
        "## 三、方案设计",
        PROBLEM_MODELING.strip() + "\n\n## 四、方案设计",
    )
    text = text.replace(
        "## 四、实验设计与结果分析",
        IMPLEMENTATION.strip() + "\n\n## 六、实验设计与结果分析",
    )
    text = text.replace("## 五、风险与安全分析", "## 七、风险与安全分析")
    text = text.replace("## 六、价值评估与可行性", "## 八、价值评估与可行性")
    text = text.replace("## 七、局限与展望", "## 九、局限与展望")
    text = text.replace("## 八、结论", "## 十、结论")
    text = text.replace(
        "DeFi 抵押借贷协议通常依赖智能合约和预言机价格实现自动化清算。",
        "DeFi 抵押借贷协议通常依赖智能合约和预言机价格实现自动化清算[1][2][5]。",
        1,
    )
    text = text.replace(
        "预言机负责将链下数据传递到链上智能合约。",
        "预言机负责将链下数据传递到链上智能合约[5][6]。",
        1,
    )
    text = text.replace(
        "预言机可能面临以下风险：",
        "预言机可能面临以下风险[7][8][9]：",
        1,
    )
    text = text.replace(
        "本文设计一个简化的 DeFi 抵押借贷清算系统，用于展示预言机价格不确定性如何影响清算决策，并提出不确定性缩放部分清算机制 USPL。",
        "本文设计一个简化的 DeFi 抵押借贷清算系统，用于展示预言机价格不确定性如何影响清算决策，并提出一种不确定性缩放清算比例机制 USPL。",
    )
    text = text.replace(
        "USPL 的核心思想是：不将预言机价格视为完全确定的单点，而是构造价格不确定区间，并让清算强度随不确定性大小动态变化。",
        "USPL 的核心思想是：不将预言机价格视为完全确定的单点，也不把边界账户简单处理为完全可清算或完全不可清算，而是构造价格不确定区间、计算健康因子区间，并让最大清算比例随不确定性大小动态变化。",
    )
    text = text.replace(
        "### 3.5 系统流程",
        DIFFERENCE_SECTION.strip() + "\n\n### 3.6 系统流程",
    )
    text = text.replace(
        "本文方案的核心机制框架如图 1 所示。与工程实现图不同，该图只保留论文分析中最关键的三层关系：预言机价格不确定性的来源、USPL 的区间化清算决策，以及最终风险结果。这样可以突出本文机制贡献，而不是展示完整智能合约调用细节。",
        "本文方案的核心机制框架如图 1 所示。与工程实现图不同，该图只保留论文分析中最关键的三层关系：预言机价格不确定性的来源、USPL 的区间化风险识别与清算比例调节，以及最终风险结果。这样可以突出本文的机制分析重点，而不是展示完整智能合约调用细节。",
    )
    text = text.replace(
        "这说明 USPL 的参数不能依赖主观设定，也不能通过反复调参追求单一场景下的最优结果。更合理的做法是在后续研究中引入参数校准机制：先定义协议侧坏账损失和用户侧误清算损失的综合目标函数，在历史校准样本中选择 `gamma`、`cap_min` 和 `cap_max`，再在独立压力场景中检验机制效果。本文课程阶段仅进行敏感性分析，不将自动调参作为已完成贡献。",
        "这说明 USPL 的参数不能依赖主观设定，也不能通过反复调参追求单一场景下的最优结果。更合理的做法是在后续研究中引入参数校准机制：先定义协议侧坏账损失和用户侧误清算损失的综合目标函数，在历史校准样本中选择 `gamma`、`cap_min` 和 `cap_max`，再在独立压力场景中检验机制效果。本文仅进行敏感性分析，不将自动调参作为已完成贡献。",
    )
    text = text.replace(
        "本文已实现课程级原型，包括：",
        "本文已实现简化实验原型，包括：",
    )
    text = text.replace(
        "该 MVP 不依赖真实主网部署，适合课程展示和机制说明。",
        "该 MVP 不依赖真实主网部署，适合进行机制展示、参数对比和风险分析。",
    )
    text = text.replace(
        "本文围绕 DeFi 抵押借贷中的预言机驱动清算问题，分析了坏账风险和误清算风险之间的机制冲突，并提出了 USPL 不确定性缩放部分清算机制。该机制将预言机价格从单点估计扩展为价格区间，并进一步计算健康因子区间，在不确定区中根据预言机不确定性动态限制最大清算比例。",
        "本文围绕 DeFi 抵押借贷中的预言机驱动清算问题，分析了坏账风险和误清算风险之间的机制冲突，并提出了 USPL 不确定性缩放清算比例机制。该机制将预言机价格从单点估计扩展为价格区间，并进一步计算健康因子区间，在边界不确定状态下根据预言机不确定性动态限制最大清算比例。",
    )
    text = text.replace(
        "总体而言，USPL 为 DeFi 清算机制提供了一种轻量、可解释、可合约化实现的改进思路，能够较好体现区块链课程中智能合约、预言机和链上自动执行机制的技术特点。",
        "总体而言，USPL 为 DeFi 清算机制提供了一种轻量、可解释、可合约化表达的改进思路，能够展示智能合约、预言机和链上自动执行机制在风险控制中的相互作用。",
    )
    text = text.replace("课程级", "实验性")
    text = text.replace("课程阶段", "本文实验阶段")
    text = text.replace("课程报告", "本文")
    text = text.replace("课程实验中", "本文实验中")
    return text.rstrip() + "\n\n" + REFERENCES.strip() + "\n"


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph()
    if level == 1:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_run_font(run, size=15, bold=True, font_cn="黑体")
    else:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, font_cn="黑体")
    return p


def add_body_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    set_paragraph_body(p, first_line=True)
    add_text_with_citations(p, text, size=12)


def add_reference_paragraph(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-21)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text.replace("**", "").replace("`", ""))
    set_run_font(run, size=10.5)


def add_keywords_paragraph(doc: Document, text: str):
    clean = text.replace("**", "").replace("`", "")
    label = "关键词："
    body = clean
    if clean.startswith(label):
        body = clean[len(label):]
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.first_line_indent = Pt(24)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label)
    set_run_font(run, size=12, bold=True)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p2.paragraph_format.first_line_indent = Pt(24)
    p2.paragraph_format.line_spacing = 1.25
    p2.paragraph_format.space_after = Pt(6)
    run2 = p2.add_run(body)
    set_run_font(run2, size=12)


def add_code_block(doc: Document, lines: list[str]):
    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Pt(24)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(line)
        run.font.name = "Courier New"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Courier New")
        run.font.size = Pt(9.5)


def add_bullet(doc: Document, text: str):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18)
    p.paragraph_format.first_line_indent = Pt(-18)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run("• ")
    set_run_font(run, size=12)
    add_text_with_citations(p, text, size=12)


def add_markdown_table(doc: Document, rows: list[str]):
    parsed = []
    for row in rows:
        cells = [c.strip() for c in row.strip().strip("|").split("|")]
        if all(re.fullmatch(r":?-{3,}:?", c) for c in cells):
            continue
        parsed.append(cells)
    if not parsed:
        return
    table = doc.add_table(rows=len(parsed), cols=len(parsed[0]))
    table.style = "Table Grid"
    for r_idx, cells in enumerate(parsed):
        for c_idx, cell_text in enumerate(cells):
            cell = table.cell(r_idx, c_idx)
            set_cell_text(
                cell,
                cell_text,
                size=10.5,
                bold=(r_idx == 0),
                align=WD_ALIGN_PARAGRAPH.CENTER if c_idx != 1 else WD_ALIGN_PARAGRAPH.LEFT,
            )
            if r_idx == 0:
                set_cell_shading(cell, "EDEDED")
    doc.add_paragraph()


def add_image(doc: Document, alt: str, path_text: str):
    path = Path(path_text)
    if not path.exists():
        add_body_paragraph(doc, f"（图像文件未找到：{path_text}）")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(5.8))


def parse_markdown_into_doc(doc: Document, text: str):
    lines = text.splitlines()
    i = 0
    in_code = False
    code_lines: list[str] = []
    table_lines: list[str] = []
    current_section_number: int | None = None
    current_sub_number = 0
    in_references = False
    cn_nums = {
        "一": 1,
        "二": 2,
        "三": 3,
        "四": 4,
        "五": 5,
        "六": 6,
        "七": 7,
        "八": 8,
        "九": 9,
        "十": 10,
    }

    def flush_table():
        nonlocal table_lines
        if table_lines:
            add_markdown_table(doc, table_lines)
            table_lines = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            flush_table()
            i += 1
            continue

        if line.startswith("|") and line.endswith("|"):
            table_lines.append(line)
            i += 1
            continue

        flush_table()

        image_match = re.match(r"!\[(.*?)\]\((.*?)\)", line)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            i += 1
            continue

        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(12)
            run = p.add_run(line[2:].strip())
            set_run_font(run, size=18, bold=True, font_cn="黑体")
            i += 1
            continue

        if line.startswith("## "):
            text_heading = line[3:].strip()
            if text_heading == "参考文献":
                doc.add_section(WD_SECTION.NEW_PAGE)
                current_section_number = None
                in_references = True
            elif text_heading.startswith("附录"):
                in_references = False
            else:
                m = re.match(r"([一二三四五六七八九十])、", text_heading)
                if m:
                    current_section_number = cn_nums[m.group(1)]
                    current_sub_number = 0
                in_references = False
            add_heading(doc, text_heading, 1)
            i += 1
            continue

        if line.startswith("### "):
            sub_heading = line[4:].strip()
            if current_section_number is not None:
                current_sub_number += 1
                sub_heading = re.sub(
                    r"^\d+\.\d+",
                    f"{current_section_number}.{current_sub_number}",
                    sub_heading,
                    count=1,
                )
            add_heading(doc, sub_heading, 2)
            i += 1
            continue

        if line.startswith("- "):
            add_bullet(doc, line[2:].strip())
            i += 1
            continue

        if line.startswith("**关键词**") or line.startswith("关键词"):
            add_keywords_paragraph(doc, line)
            i += 1
            continue

        if line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(24)
            p.paragraph_format.right_indent = Pt(12)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            add_text_with_citations(p, line[2:].strip(), size=12, bold=True, font_cn="楷体")
            i += 1
            continue

        # Merge adjacent plain lines into one paragraph.
        parts = [line.strip()]
        j = i + 1
        while j < len(lines):
            nxt = lines[j].strip()
            if (
                not nxt
                or nxt.startswith("#")
                or nxt.startswith("- ")
                or nxt.startswith("|")
                or nxt.startswith("```")
                or nxt.startswith("![")
                or nxt.startswith("> ")
            ):
                break
            parts.append(nxt)
            j += 1
        if in_references:
            add_reference_paragraph(doc, " ".join(parts))
        else:
            add_body_paragraph(doc, " ".join(parts))
        i = j

    flush_table()


def set_document_defaults(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(12)
    for section in doc.sections:
        section.page_width = Cm(21)
        section.page_height = Cm(29.7)
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.17)
        section.right_margin = Cm(3.17)


def main():
    doc = Document()
    set_document_defaults(doc)
    build_cover(doc)

    text = prepare_markdown()
    parse_markdown_into_doc(doc, text)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
